from __future__ import annotations

import sys

import pytest
from docx import Document
from pptx import Presentation


def _load_server(tmp_path, monkeypatch):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    monkeypatch.setattr(sys, "argv", ["crisai-test-document-export", str(workspace)])
    for name in list(sys.modules):
        if name == "crisai.servers.document_export_server" or name.startswith("crisai.servers.document_export_server."):
            del sys.modules[name]
    import crisai.servers.document_export_server as document_export_server

    return document_export_server, workspace


def test_render_docx_from_markdown_uses_manifest_and_reports_missing_sections(tmp_path, monkeypatch):
    server, workspace = _load_server(tmp_path, monkeypatch)
    source = workspace / "tasks/demo/artefacts/design.md"
    manifest = workspace / "knowledge/templates/ucl/hld/ucl-hld-docx.template.yaml"
    output = workspace / "tasks/demo/exports/design.docx"
    source.parent.mkdir(parents=True)
    manifest.parent.mkdir(parents=True)
    source.write_text(
        "---\ntitle: Reporting HLD\n---\n\n"
        "## Context\nA reporting design.\n\n"
        "## Target architecture\n- Curated dataset\n",
        encoding="utf-8",
    )
    manifest.write_text(
        "template_id: test_hld\n"
        "output_type: docx\n"
        "required_sections:\n"
        "  - Context\n"
        "  - Target architecture\n"
        "  - Key decisions\n",
        encoding="utf-8",
    )

    report = server.render_docx_from_markdown(
        "tasks/demo/artefacts/design.md",
        "knowledge/templates/ucl/hld/ucl-hld-docx.template.yaml",
        "tasks/demo/exports/design.docx",
    )

    assert report["output_path"] == "tasks/demo/exports/design.docx"
    assert "Missing required section: Key decisions" in report["warnings"]
    doc = Document(str(output))
    assert any("Reporting HLD" in para.text for para in doc.paragraphs)
    assert output.exists()


def test_render_pptx_from_markdown_creates_deck(tmp_path, monkeypatch):
    server, workspace = _load_server(tmp_path, monkeypatch)
    source = workspace / "tasks/demo/artefacts/design.md"
    manifest = workspace / "knowledge/templates/ucl/presentation/ucl-pptx.template.yaml"
    output = workspace / "tasks/demo/exports/design.pptx"
    source.parent.mkdir(parents=True)
    manifest.parent.mkdir(parents=True)
    source.write_text(
        "---\ntitle: Architecture Summary\n---\n\n"
        "## Context\n- Current reporting is manual.\n\n"
        "## Key decisions\n- Use curated semantic model.\n",
        encoding="utf-8",
    )
    manifest.write_text(
        "template_id: test_pptx\n"
        "template_name: Test Deck\n"
        "output_type: pptx\n"
        "required_sections:\n"
        "  - Context\n"
        "  - Key decisions\n",
        encoding="utf-8",
    )

    report = server.render_pptx_from_markdown(
        "tasks/demo/artefacts/design.md",
        "knowledge/templates/ucl/presentation/ucl-pptx.template.yaml",
        "tasks/demo/exports/design.pptx",
    )

    deck = Presentation(str(output))
    assert report["warnings"] == []
    assert len(deck.slides) == 3
    assert output.exists()


def test_inspect_document_template_manifest_returns_contract(tmp_path, monkeypatch):
    server, workspace = _load_server(tmp_path, monkeypatch)
    manifest = workspace / "knowledge/templates/ucl/hld/ucl-hld-docx.template.yaml"
    manifest.parent.mkdir(parents=True)
    manifest.write_text(
        "template_id: test_hld\n"
        "template_name: Test HLD\n"
        "output_type: docx\n"
        "description: Test template manifest.\n"
        "required_sections:\n"
        "  - Context\n"
        "  - Target architecture\n",
        encoding="utf-8",
    )

    report = server.inspect_document_template_manifest(
        "knowledge/templates/ucl/hld/ucl-hld-docx.template.yaml"
    )

    assert report == {
        "path": "knowledge/templates/ucl/hld/ucl-hld-docx.template.yaml",
        "template_id": "test_hld",
        "template_name": "Test HLD",
        "output_type": "docx",
        "required_sections": ["Context", "Target architecture"],
        "template_file": None,
        "description": "Test template manifest.",
    }


def test_render_export_rejects_non_export_output_path(tmp_path, monkeypatch):
    server, workspace = _load_server(tmp_path, monkeypatch)
    source = workspace / "tasks/demo/artefacts/design.md"
    manifest = workspace / "knowledge/templates/ucl/hld/ucl-hld-docx.template.yaml"
    source.parent.mkdir(parents=True)
    manifest.parent.mkdir(parents=True)
    source.write_text("## Context\nx\n", encoding="utf-8")
    manifest.write_text("template_id: test\noutput_type: docx\n", encoding="utf-8")

    try:
        server.render_docx_from_markdown(
            "tasks/demo/artefacts/design.md",
            "knowledge/templates/ucl/hld/ucl-hld-docx.template.yaml",
            "tasks/demo/artefacts/design.docx",
        )
    except ValueError as exc:
        assert "tasks/<task>/exports" in str(exc)
    else:
        raise AssertionError("expected output path validation failure")


def test_inspect_document_template_manifest_blocks_sensitive_secret_folder(tmp_path, monkeypatch):
    server, workspace = _load_server(tmp_path, monkeypatch)
    manifest = workspace / ".secrets" / "template.yaml"
    manifest.parent.mkdir()
    manifest.write_text("template_id: secret\n", encoding="utf-8")

    with pytest.raises(ValueError, match="restricted"):
        server.inspect_document_template_manifest(".secrets/template.yaml")


def test_render_docx_blocks_sensitive_manifest_template_file(tmp_path, monkeypatch):
    server, workspace = _load_server(tmp_path, monkeypatch)
    source = workspace / "tasks/demo/artefacts/design.md"
    manifest = workspace / "knowledge/templates/ucl/hld/ucl-hld-docx.template.yaml"
    secret_template = workspace / ".auth" / "template.docx"
    output = workspace / "tasks/demo/exports/design.docx"
    source.parent.mkdir(parents=True)
    manifest.parent.mkdir(parents=True)
    secret_template.parent.mkdir()
    source.write_text("## Context\nx\n", encoding="utf-8")
    manifest.write_text(
        "template_id: test\n"
        "output_type: docx\n"
        "template_file: ../../../../.auth/template.docx\n",
        encoding="utf-8",
    )
    Document().save(str(secret_template))

    with pytest.raises(ValueError, match="restricted"):
        server.render_docx_from_markdown(
            "tasks/demo/artefacts/design.md",
            "knowledge/templates/ucl/hld/ucl-hld-docx.template.yaml",
            "tasks/demo/exports/design.docx",
        )
    assert not output.exists()
